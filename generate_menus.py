#!/usr/bin/env python3
"""
Henbrook Daily Menu Generator (refactor)

What this tool does
-------------------
1) Parse a weekly DOCX grid into structured "day" data.
2) Render Standard + Vegan daily menus (docxtpl) and the Allergens sheet (python-docx).
3) Write **only a ZIP per day** containing:
   - Residents_DD-MM-YYYY.docx
   - Residents_DD-MM-YYYY_vegan.docx
   - Allergens_Residents_DD-MM-YYYY.docx

Key rules preserved
-------------------
• Do not alter template styles; populate placeholders only.
• Sentence-case dessert titles; (V) for Standard desserts, (Ve) for Vegan desserts.
• Vegan lunch includes Jacket potato (short title); fixed allergens ticks in allergens table.
• Vegan allergens scrub Milk/Eggs.
• Standard/vegan soups & dessert highlighting per spec.
• Allergens table:
    - Section rows “— Standard —”, “— Vegan —”
    - Vegan lunch starter = actual weekly title + “(Ve)”
    - Only **one** vegan supper soup line
    - Order mirrors Standard section: (Lunch starter[s]) → mains → optional sides → Lunch dessert → Supper soup → Supper special → Supper dessert.
• Date banner in allergens sheet becomes “D Month YYYY”.

CLI
---
Either provide a templates folder (default ./templates) or pass template paths.
See bottom of file for CLI examples.
"""

from __future__ import annotations

# ──────────────────────────────────────────────────────────────────────────────
# Imports
# ──────────────────────────────────────────────────────────────────────────────
import argparse
import io
import os
import re
import zipfile
import datetime as dt
from typing import Dict, List, Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docxtpl import DocxTemplate, RichText

# ──────────────────────────────────────────────────────────────────────────────
# Small utilities: formatting, text, highlighting
# ──────────────────────────────────────────────────────────────────────────────

def _yellow(text: str) -> RichText:
    """Return the given text as a yellow-highlight RichText run."""
    rt = RichText()
    rt.add(text or "", highlight="yellow")
    return rt

def _is_chefs_choice_soup(title: str) -> bool:
    s = (str(title) or "").lower().replace("’", "'")
    return ("chef" in s) and ("choice" in s) and ("soup" in s)

def sentence_case(s: str) -> str:
    s = (s or "").strip()
    if not s:
        return s
    first_alpha_idx = next((i for i, c in enumerate(s) if c.isalpha()), None)
    if first_alpha_idx is None:
        return s
    return s[:first_alpha_idx] + s[first_alpha_idx].upper() + s[first_alpha_idx + 1 :].lower()

def strip_suffixes(s: str) -> str:
    return re.sub(r"\s*\((Ve|V)\)\s*$", "", s or "").strip()

def add_suffix(s: str, suffix: str) -> str:
    return f"{strip_suffixes(s)} {suffix}".strip()

def normalise_sides(text: str) -> str:
    """Turn messy weekly 'Optional Sides' content into 'A, B, C' (de-duplicated)."""
    if not text:
        return ""
    raw = text.replace("\r", "\n")
    parts = re.split(r",|/|\t+|\s{2,}|\n+", raw)
    parts = [p.strip(" -–•").rstrip(".").strip() for p in parts if p and p.strip()]
    seen, uniq = set(), []
    for p in parts:
        if p.lower() not in seen:
            seen.add(p.lower())
            uniq.append(p)
    return ", ".join(uniq)

def ensure_len(lst: List[Any], n: int, filler=None) -> List[Any]:
    out = list(lst)
    while len(out) < n:
        out.append(filler if filler is not None else {})
    return out[:n]

def to_iso(date_str: str) -> str:
    ds = date_str.strip()
    for fmt in ("%d/%m/%Y", "%d/%m/%y", "%Y-%m-%d"):
        try:
            d = dt.datetime.strptime(ds, fmt).date()
            return d.isoformat()
        except Exception:
            pass
    return ds

def date_label(iso_date: str, weekday_name: str) -> str:
    """Weekday – dd/mm/yyyy as used by your templates."""
    try:
        d = dt.date.fromisoformat(iso_date)
        return f"{weekday_name} – {d.strftime('%d/%m/%Y')}"
    except Exception:
        return f"{weekday_name} – {iso_date}"

def pretty_month_banner(d: dt.date) -> str:
    """Allergens header banner, e.g. '18 September 2025' (no ordinal suffix)."""
    return f"{d.day} {d.strftime('%B %Y')}"

# ──────────────────────────────────────────────────────────────────────────────
# Allergens mapping / parsing
# ──────────────────────────────────────────────────────────────────────────────

CANON_COLS = [
    "Celery",
    "Cereals with Gluten",
    "Crustaceans",
    "Eggs",
    "Fish",
    "Lupin",
    "Milk",
    "Molluscs",
    "Mustards",
    "Peanuts",
    "Nuts from Trees",
    "Sesame",
    "Soybeans",
    "Sulphur",
    "Alcohol",
    "Pork",
]

ALLERGEN_NORMALISE = {
    "celery": "Celery",
    "gluten": "Cereals with Gluten",
    "cereals with gluten": "Cereals with Gluten",
    "crustaceans": "Crustaceans",
    "egg": "Eggs",
    "eggs": "Eggs",
    "fish": "Fish",
    "lupin": "Lupin",
    "milk": "Milk",
    "mollusc": "Molluscs",
    "molluscs": "Molluscs",
    "mustard": "Mustards",
    "mustards": "Mustards",
    "peanut": "Peanuts",
    "peanuts": "Peanuts",
    "tree nuts": "Nuts from Trees",
    "nuts": "Nuts from Trees",
    "sesame": "Sesame",
    "soya": "Soybeans",
    "soy": "Soybeans",
    "soybeans": "Soybeans",
    "sulphite": "Sulphur",
    "sulphites": "Sulphur",
    "sulfur dioxide": "Sulphur",
    "sulphur dioxide": "Sulphur",
    "sulphur": "Sulphur",
    "alcohol": "Alcohol",
    "pork": "Pork",
}

NON_VEGAN_CANON = {"Milk", "Eggs"}  # scrub from vegan items

ALLERGEN_TAIL_RE = re.compile(
    r"\b(?:Milk|Egg|Soya|Gluten|Nuts|Peanuts|Sesame|Mustard|Celery|Sulphites|Fish|Crustaceans|Molluscs)\b.*$",
    re.IGNORECASE,
)

def _canon_from_header(text: str) -> str | None:
    t = (text or "").lower()
    t = t.replace("doi", "dio").replace("d02", "dio")
    checks = [
        ("cereal" in t and "gluten" in t, "Cereals with Gluten"),
        ("celery" in t, "Celery"),
        ("crustace" in t, "Crustaceans"),
        ("egg" in t, "Eggs"),
        ("fish" in t, "Fish"),
        ("lupin" in t, "Lupin"),
        ("milk" in t, "Milk"),
        ("mollusc" in t, "Molluscs"),
        ("mustard" in t, "Mustards"),
        ("peanut" in t, "Peanuts"),
        (("nut" in t and "tree" in t) or "nuts from trees" in t, "Nuts from Trees"),
        ("sesame" in t, "Sesame"),
        ("soya" in t or "soy" in t, "Soybeans"),
        ("sulph" in t or "sulfur" in t, "Sulphur"),
        ("alcohol" in t, "Alcohol"),
        ("pork" in t, "Pork"),
    ]
    for cond, name in checks:
        if cond:
            return name
    return None

def _parse_allergen_csv(csv_text: str) -> set[str]:
    """CSV or slash-separated string → set of canonical column names."""
    if not csv_text:
        return set()
    raw = csv_text.replace("/", ",")
    toks = [t.strip().lower() for t in raw.split(",") if t.strip()]
    return {ALLERGEN_NORMALISE[t] for t in toks if t in ALLERGEN_NORMALISE}

def _scrub_vegan_csv_to_canonset(csv_text: str) -> set[str]:
    return {c for c in _parse_allergen_csv(csv_text) if c not in NON_VEGAN_CANON}

def _remove_tokens_from_csv(csv_text: str, forbidden={"milk", "egg", "eggs"}) -> str:
    """Remove forbidden tokens from a CSV/slash list (string form for menus)."""
    if not csv_text:
        return ""
    raw = csv_text.replace("/", ",")
    toks = [t.strip() for t in raw.split(",") if t.strip()]
    keep = [t for t in toks if t.lower() not in forbidden]
    return ", ".join(keep)

# ──────────────────────────────────────────────────────────────────────────────
# Weekly grid parsing
# ──────────────────────────────────────────────────────────────────────────────

def clean_text(s: str) -> str:
    if s is None:
        return ""
    s = s.replace("\r", "")
    s = re.sub(r"[ \t]+", " ", s)
    lines = [ln.strip() for ln in s.split("\n")]
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()
    return "\n".join(lines)

def parse_title_desc_allergens(cell_text: str) -> Dict[str, str]:
    """
    Robustly split a weekly cell into:
      title, description (optional), allergens (optional).
    Accepts 'Title\\nDescription\\nAllergens: ...' or same-line allergens.
    """
    raw = (cell_text or "").replace("\r", "\n").strip()

    # peel off allergens 'tail'
    allergens = ""
    m = ALLERGEN_TAIL_RE.search(raw)
    if m:
        allergens = raw[m.start():].strip()
        body = raw[:m.start()].strip()
    else:
        body = raw

    # split the body; normalize long spaces to newlines
    body = re.sub(r"[ \t]{2,}", "\n", body)
    lines = [ln.strip() for ln in body.split("\n") if ln.strip()]

    title = lines[0] if lines else ""
    description = " ".join(lines[1:]).strip() if len(lines) > 1 else ""

    if allergens:
        allergens = re.sub(r"\s*,\s*", ", ", allergens)
        allergens = allergens.replace(" ,", ",").strip(" ,")

    return {"title": title, "description": description, "allergens": allergens}

def parse_week(weekly_docx_path: str) -> Dict[str, Any]:
    """Read the weekly DOCX grid and return {'days': [ ...7 day dicts... ]}."""
    doc = Document(weekly_docx_path)
    tbl = doc.tables[0]
    rows = tbl.rows

    def cell(r, c):
        try:
            return rows[r].cells[c].text
        except Exception:
            return ""

    # row/column map (0-based)
    # r0: weekday names; r1: dates; r2: themes; r3..10 lunch; r12..17 supper
    day_names = [cell(0, c + 1).strip() for c in range(7)]
    dates_iso = [to_iso(cell(1, c + 1).strip()) for c in range(7)]
    themes = [cell(2, c + 1).strip() for c in range(7)]

    # BEFORE
    interesting_rows = [3,4,5,6,7,8,9,10,12,15,16,17]

    # AFTER  (add 13 and 14 so supper vegan + selection are available)
    interesting_rows = [3,4,5,6,7,8,9,10,12,13,14,15,16,17]
    per_row_per_day: Dict[int, Dict[int, Dict[str, str]]] = {r: {} for r in interesting_rows}
    for r in interesting_rows:
        for c in range(7):
            per_row_per_day[r][c] = parse_title_desc_allergens(cell(r, c + 1))

    def daily(day_idx: int) -> Dict[str, Any]:
        day_name = day_names[day_idx] or ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"][day_idx]
        date_iso = dates_iso[day_idx]
        theme = themes[day_idx]
        g = {r: per_row_per_day[r][day_idx] for r in interesting_rows}

        lunch_starters = [
            {"title": g[3].get("title", ""), "allergens": g[3].get("allergens", "")},
            {"title": g[4].get("title", ""), "allergens": g[4].get("allergens", "")},
        ]
        veg_main = g[6]
        meat_main = g[7]
        vegan_main = g[5]
        opt_raw = " ".join([g[8].get("title", "") or "", g[8].get("description", "") or ""]).strip()
        opt_sides = {"title": normalise_sides(opt_raw), "allergens": g[8].get("allergens", "")}
        desserts = [
            {"title": g[9].get("title", ""), "allergens": g[9].get("allergens", "")},
            {"title": "Ice creams / sorbet (V)", "allergens": g[10].get("allergens", "")},
        ]
        supper_special = g[15]
        vegan_supper_special = g[13]  # vegan supper special (row r14)
        supper_desserts = [
            {"title": g[16].get("title", ""), "allergens": g[16].get("allergens", "")},
            {
                "title": "Selection of vegan ice creams or sorbet with seasonal fruits (Ve)",
                "allergens": g[17].get("allergens", ""),
            },
        ]
        return {
            "header": {
                "theme": theme,
                "date": date_label(date_iso, day_name),
                "date_iso": date_iso,
                "weekday": day_name,
            },
            "source": {
                "lunch": {
                    "starters": ensure_len(lunch_starters, 2, {"title": "", "allergens": ""}),
                    "veg_main": veg_main,
                    "meat_main": meat_main,
                    "vegan_main": vegan_main,
                    "optional_sides": opt_sides,
                    "desserts": ensure_len(desserts, 2, {"title": "", "allergens": ""}),
                },
                "supper": {
                    "starter":   {"title": "Chef’s choice soup", "allergens": g[12].get("allergens", "")},
                    "selection": {"allergens": "Sulphites, Gluten, Mustard, Soya"},
                    "specials":  supper_special,          # standard
                    "vegan_special": vegan_supper_special,  # ← add this
                    "desserts":  ensure_len(supper_desserts, 2, {"title":"", "allergens":""}),
                },
            },
        }

    return {"days": [daily(i) for i in range(7)]}

# ──────────────────────────────────────────────────────────────────────────────
# Build docxtpl contexts
# ──────────────────────────────────────────────────────────────────────────────

def build_standard_context(day: Dict[str, Any]) -> Dict[str, Any]:
    src = day["source"]; header = day["header"]
    starters = ensure_len(src["lunch"]["starters"], 2, {"title": "", "allergens": ""})
    mv, mm = src["lunch"]["veg_main"], src["lunch"]["meat_main"]

    mains = [
        {"title": strip_suffixes(mv.get("title", "")) + " (V)", "description": mv.get("description", ""), "allergens": mv.get("allergens", "")},
        {"title": strip_suffixes(mm.get("title", "")), "description": mm.get("description", ""), "allergens": mm.get("allergens", "")},
    ]
    opt = src["lunch"]["optional_sides"]
    d0_title = sentence_case(strip_suffixes(src["lunch"]["desserts"][0]["title"])) + " (V)"
    d1_all = src["lunch"]["desserts"][1]["allergens"]
    supper_starter = {"title": "Chef’s choice soup (V)", "allergens": src["lunch"]["starters"][0].get("allergens", "")}
    sp = src["supper"]["specials"]
    sd0_title = sentence_case(strip_suffixes(src["supper"]["desserts"][0]["title"])) + " (V)"
    sd1_all = src["supper"]["desserts"][1]["allergens"]

    # (Optional) highlight soups in Standard too (if you kept this earlier)
    if _is_chefs_choice_soup(starters[0]["title"]): starters[0]["title"] = _yellow(starters[0]["title"])
    if _is_chefs_choice_soup(starters[1]["title"]): starters[1]["title"] = _yellow(starters[1]["title"])
    if _is_chefs_choice_soup(supper_starter["title"]): supper_starter["title"] = _yellow(supper_starter["title"])

    return {
        "header": {"theme": header["theme"], "date": header["date"]},
        "days":   {"theme": header["theme"], "date": header["date"]},
        "lunch": {
            "starters": [{"title": starters[0]["title"], "allergens": starters[0]["allergens"]},
                         {"title": starters[1]["title"], "allergens": starters[1]["allergens"]}],
            "mains": mains,
            "optional_sides": {"title": opt.get("title", ""), "allergens": opt.get("allergens", "")},
            "desserts": [
                {"title": d0_title, "allergens": src["lunch"]["desserts"][0]["allergens"]},
                {"title": src["lunch"]["desserts"][1]["title"], "allergens": d1_all},
            ],
        },
        "supper": {
            "starter": supper_starter,
            "selection": {"allergens": "Sulphites, Gluten, Mustard, Soya"},
            "specials": {"title": strip_suffixes(sp.get("title","")), "description": sp.get("description",""), "allergens": sp.get("allergens","")},
            "desserts": [
                {"title": sd0_title, "allergens": src["supper"]["desserts"][0]["allergens"]},
                {"title": src["supper"]["desserts"][1]["title"], "allergens": sd1_all},
            ],
        },
    }

def build_vegan_context(day: Dict[str, Any]) -> Dict[str, Any]:
    src    = day["source"]
    header = day["header"]

    # --- LUNCH starters -------------------------------------------------------
    starters = ensure_len(src["lunch"]["starters"], 2, {"title": "", "allergens": ""})
    lunch_soup_title = add_suffix(strip_suffixes(starters[0].get("title", "")), "(Ve)")   # not highlighted
    lunch_soup_all   = _remove_tokens_from_csv(starters[0].get("allergens", ""))         # scrub Milk/Egg

    # --- LUNCH mains (vegan) --------------------------------------------------
    vmain  = src["lunch"]["vegan_main"]   # weekly vegan main row (r5)
    vegstd = src["lunch"]["veg_main"]     # Standard vegetarian row (r6) — for fallback desc

    # Title/allergens from vegan row, guard against title dumped in allergens
    v_title_raw = vmain.get("title", "") or ""
    v_all_raw   = vmain.get("allergens", "") or ""
    if "(ve" in v_all_raw.lower() or re.search(r"[A-Za-z].+\(ve\)", v_all_raw, re.I):
        m = re.search(r"([A-Za-z].*?\(ve\))", v_all_raw, flags=re.I)
        if m and len(strip_suffixes(v_title_raw)) < 2:
            v_title_raw = m.group(1).strip()
        v_all_raw = re.sub(r"[A-Za-z].*?\(ve\)\s*", "", v_all_raw, flags=re.I).strip()

    vegan_main_title = add_suffix(strip_suffixes(v_title_raw), "(Ve)")
    vegan_main_desc  = (vmain.get("description", "") or "").strip()
    if not vegan_main_desc:
        borrowed = (vegstd.get("description", "") or "").strip()
        vegan_main_desc = _yellow(borrowed) if borrowed else ""     # highlight borrowed text
    vegan_main_all   = _remove_tokens_from_csv(v_all_raw)           # scrub Milk/Egg from menu display

    # Optional sides (single block)
    opt_title = src["lunch"]["optional_sides"].get("title", "")
    opt_all   = _remove_tokens_from_csv(src["lunch"]["optional_sides"].get("allergens", ""))

    # Vegan desserts (fixed allergens; highlight dessert-1 only)
    vegan_desserts_all = "Gluten, Nuts, Soya, Sulphites"
    lunch_dessert_title = _yellow(sentence_case(strip_suffixes(src["lunch"]["desserts"][0]["title"])) + " (Ve)")
    lunch_ice_title     = "Selection of vegan ice creams or sorbet with seasonal fruits (Ve)"

    # --- SUPPER starter -------------------------------------------------------
    # Title highlighted; allergens MUST COPY lunch soup allergens
    supper_soup_title = _yellow("Chef’s choice soup (Ve)")
    supper_soup_all   = _remove_tokens_from_csv(starters[0].get("allergens", ""))

    # --- SUPPER special (vegan) — apply SAME safety net as lunch main ---------
    sp_std = src["supper"]["specials"]                 # standard row (r16)
    sp_veg = src["supper"].get("vegan_special", {})    # vegan row (r14)

    vs_title_raw = sp_veg.get("title", "") or ""
    vs_all_raw   = sp_veg.get("allergens", "") or ""

    # SAFETY NET: if weekly cell dumped the '(Ve)' dish title into the allergens line, recover it
    if "(ve" in vs_all_raw.lower() or re.search(r"[A-Za-z].+\(ve\)", vs_all_raw, re.I):
        m = re.search(r"([A-Za-z].*?\(ve\))", vs_all_raw, flags=re.I)
        if m and len(strip_suffixes(vs_title_raw)) < 2:
            vs_title_raw = m.group(1).strip()
        vs_all_raw = re.sub(r"[A-Za-z].*?\(ve\)\s*", "", vs_all_raw, flags=re.I).strip()

    # If vegan title still empty, fall back to extracting a vegan variant from the standard title/desc
    if not vs_title_raw.strip():
        vs_title_raw = add_suffix(strip_suffixes(
            pick_vegan_variant(sp_std.get("title",""), sp_std.get("description",""))
        ), "(Ve)")

    veg_special_title = add_suffix(strip_suffixes(vs_title_raw), "(Ve)")
    veg_special_desc  = (sp_veg.get("description", "") or "").strip()
    if not veg_special_desc:
        borrowed = (sp_std.get("description", "") or "").strip()
        veg_special_desc = _yellow(borrowed) if borrowed else ""     # highlight borrowed supper description
    veg_special_all   = _remove_tokens_from_csv(vs_all_raw)

    # SUPPER desserts
    supper_dessert_title = _yellow(sentence_case(strip_suffixes(src["supper"]["desserts"][0]["title"])) + " (Ve)")
    supper_ice_title     = "Selection of vegan ice creams or sorbet with seasonal fruits (Ve)"

    # Assemble context matching the template schema (scalar placeholders only)
    return {
        "header": {"theme": header["theme"], "date": header["date"]},
        "days":   {"theme": header["theme"], "date": header["date"]},  # for {{days.*}} placeholders
        "lunch": {
            "starters": [
                {"title": lunch_soup_title, "allergens": lunch_soup_all},
                {"title": "", "allergens": ""},  # second starter intentionally blank on vegan menu
            ],
            "mains": [
                {"title": "Jacket potato and toppings (Ve)", "description": "",            "allergens": ""},            # jacket first
                {"title": vegan_main_title,                 "description": vegan_main_desc, "allergens": vegan_main_all},
            ],
            "optional_sides": {"title": opt_title, "allergens": opt_all},
            "desserts": [
                {"title": lunch_dessert_title, "allergens": vegan_desserts_all},  # highlighted
                {"title": lunch_ice_title,     "allergens": vegan_desserts_all},  # not highlighted
            ],
        },
        "supper": {
            "starter":   {"title": supper_soup_title, "allergens": supper_soup_all},
            "selection": {"title": "Henbrook’s assorted sandwiches (Ve)", "allergens": "Sulphites, Gluten, Mustard, Soya"},
            "specials":  {"title": veg_special_title, "description": veg_special_desc, "allergens": veg_special_all},
            "desserts": [
                {"title": supper_dessert_title, "allergens": vegan_desserts_all},  # highlighted
                {"title": supper_ice_title,     "allergens": vegan_desserts_all},  # not highlighted
            ],
        },
    }

# ──────────────────────────────────────────────────────────────────────────────
# Allergens table: build rows and write ticks
# ──────────────────────────────────────────────────────────────────────────────

def _find_allergen_table(doc: Document):
    """Return (table, header_row_idx, title_col_idx, col_map)."""
    best = None
    for tbl in doc.tables:
        for r_idx in range(min(12, len(tbl.rows))):
            row = tbl.rows[r_idx]
            col_map = {}
            for c_idx, cell in enumerate(row.cells):
                canon = _canon_from_header((cell.text or "").strip())
                if canon:
                    col_map[c_idx] = canon
            if len(col_map) >= 6:
                non_allergen_cols = [i for i in range(len(row.cells)) if i not in col_map]
                title_col = non_allergen_cols[0] if non_allergen_cols else 0
                best = (tbl, r_idx, title_col, col_map)
                break
        if best:
            break
    if not best:
        raise RuntimeError("Allergens table not found (need recognizable allergen headers).")
    return best

def _clear_cell(cell): cell.text = ""

def _set_text(cell, text: str, center=False, bold=False):
    _clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    if bold: run.bold = True

def _set_tick(cell):
    _clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("✔")
    run.font.size = Pt(12)

def _collect_allergen_items_for_table(day: dict, std_ctx: dict, veg_ctx: dict) -> list[tuple[str, set[str], bool]]:
    """
    Build [(title, set_of_columns, is_heading)] in the display order.
    - Vegan items scrub Milk/Egg.
    - Jacket potato has fixed ticks.
    - Vegan lunch starter uses actual weekly title + (Ve).
    - Only **one** vegan supper soup line.
    """
    items: list[tuple[str, set[str], bool]] = []

    def add(title, allergens, *, heading=False, force_vegan_dessert=False):
        if heading:
            items.append((str(title), set(), True)); return
        if title is None or str(title).strip() == "":
            return
        if isinstance(allergens, set):
            cols = allergens
        else:
            txt = "Gluten, Nuts, Soya, Sulphites" if force_vegan_dessert else (allergens or "")
            cols = _parse_allergen_csv(txt)
        items.append((str(title).strip(), cols, False))

    src = day["source"]

    # ----- Standard -----
    add("— Standard —", "", heading=True)
    for st in src["lunch"]["starters"]:
        add(st.get("title",""), st.get("allergens",""))
    add(src["lunch"]["veg_main"].get("title",""),  src["lunch"]["veg_main"].get("allergens",""))
    add(src["lunch"]["meat_main"].get("title",""), src["lunch"]["meat_main"].get("allergens",""))
    add(src["lunch"]["optional_sides"].get("title",""), src["lunch"]["optional_sides"].get("allergens",""))
    add(std_ctx["lunch"]["desserts"][0]["title"], std_ctx["lunch"]["desserts"][0]["allergens"])
    add("Chef’s choice soup (V)", src["supper"]["starter"]["allergens"])
    add(std_ctx["supper"]["specials"]["title"], std_ctx["supper"]["specials"]["allergens"])
    add(std_ctx["supper"]["desserts"][0]["title"], std_ctx["supper"]["desserts"][0]["allergens"])

    # ----- Vegan -----
    add("— Vegan —", "", heading=True)

    # 1) LUNCH starter (actual weekly title + (Ve)), scrubbed
    add(add_suffix(src["lunch"]["starters"][0].get("title",""), "(Ve)"),
        _scrub_vegan_csv_to_canonset(src["lunch"]["starters"][0].get("allergens","")))

    # 2) LUNCH mains (Jacket first, with fixed ticks), then vegan weekly main (scrubbed)
    for m in veg_ctx["lunch"]["mains"]:
        title = str(m["title"]); allergens = m.get("allergens", "")
        if "jacket" in title.lower():
            add("Jacket potato and toppings (Ve)", {"Celery","Sulphur","Cereals with Gluten","Mustards"})
        else:
            add(title, _scrub_vegan_csv_to_canonset(allergens))

    # 3) Optional sides (scrubbed)
    add(src["lunch"]["optional_sides"].get("title",""),
        _scrub_vegan_csv_to_canonset(src["lunch"]["optional_sides"].get("allergens","")))

    # 4) Lunch dessert 1 (fixed vegan)
    add(sentence_case(strip_suffixes(src["lunch"]["desserts"][0]["title"])) + " (Ve)",
        "Gluten, Nuts, Soya, Sulphites", force_vegan_dessert=True)

    # 5) Supper soup (Ve) – only once (use supper soup allergens, scrubbed)
    add("Chef’s choice soup (Ve)", _scrub_vegan_csv_to_canonset(src["supper"]["starter"]["allergens"]))

    # 6) SUPPER special — use vegan row (r14), scrubbed
    add(
        add_suffix(strip_suffixes(src["supper"]["vegan_special"].get("title","")), "(Ve)"),
        _scrub_vegan_csv_to_canonset(src["supper"]["vegan_special"].get("allergens","")),
    )
    # 7) Supper dessert 1 (fixed vegan)
    add(sentence_case(strip_suffixes(src["supper"]["desserts"][0]["title"])) + " (Ve)",
        "Gluten, Nuts, Soya, Sulphites", force_vegan_dessert=True)

    return items

def render_allergens_doc(day: dict, std_ctx: dict, veg_ctx: dict, allergens_tpl_path: str) -> Document:
    """Open allergens template, replace month banner, fill table, return Document."""
    doc = Document(allergens_tpl_path)
    # date banner
    d = dt.date.fromisoformat(day["header"]["date_iso"])
    banner = pretty_month_banner(d)

    def _replace_in_paragraph(p, old: str, new: str):
        if old not in p.text: return
        text = p.text.replace(old, new)
        if not p.runs: p.add_run(text); return
        p.runs[0].text = text
        for r in p.runs[1:]: r.text = ""

    for p in doc.paragraphs: _replace_in_paragraph(p, "September 2025", banner)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, "September 2025", banner)

    tbl, header_row_idx, title_col, col_map = _find_allergen_table(doc)
    data_row_start = header_row_idx + 1
    rows = _collect_allergen_items_for_table(day, std_ctx, veg_ctx)

    # write rows (leave unused rows blank)
    max_rows = len(tbl.rows) - data_row_start
    for i in range(max_rows):
        r = tbl.rows[data_row_start + i]
        for c in r.cells: _clear_cell(c)
        if i >= len(rows): continue
        title, cols, is_heading = rows[i]
        if is_heading:
            _set_text(r.cells[title_col], title, center=True, bold=True); continue
        _set_text(r.cells[title_col], title)
        for c_idx, templ_name in col_map.items():
            if templ_name in cols:
                _set_tick(r.cells[c_idx])

    return doc

# ──────────────────────────────────────────────────────────────────────────────
# Rendering helpers (to BYTES; no stray files)
# ──────────────────────────────────────────────────────────────────────────────
def pick_vegan_variant(title: str, description: str = "") -> str:
    """
    From a cell that may contain multiple variants (e.g. split by '/', newlines, 'or'),
    return the one that looks vegan (contains '(Ve)' or the word 'vegan').
    Falls back to the given title + ' (Ve)' if no explicit vegan variant is found.
    """
    t = (title or "")
    d = (description or "")
    combined = " ".join([t, d]).strip()
    if not combined:
        return add_suffix(t, "(Ve)")
    parts = re.split(r"\s*(?:/|\n|;|\||\bor\b)\s*", combined, flags=re.I)
    for p in parts:
        pl = p.lower()
        if "(ve)" in pl or "vegan" in pl:
            return add_suffix(strip_suffixes(p.strip()), "(Ve)")
    # If the original already had (Ve), keep it normalized; otherwise append.
    if "(ve)" in t.lower():
        return add_suffix(strip_suffixes(t), "(Ve)")
    return add_suffix(strip_suffixes(t), "(Ve)")

def render_docxtpl_to_bytes(template_path: str, context: dict) -> bytes:
    tpl = DocxTemplate(template_path)
    tpl.render(context)
    buf = io.BytesIO()
    tpl.docx.save(buf)
    return buf.getvalue()

def save_doc_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def render_day_to_zip(day: Dict[str, Any], standard_tpl: str, vegan_tpl: str, allergens_tpl: str, out_dir: str) -> str:
    """Render all 3 docs to memory and write a day ZIP only; return ZIP path."""
    os.makedirs(out_dir, exist_ok=True)
    std_ctx = build_standard_context(day)
    veg_ctx = build_vegan_context(day)

    d = dt.date.fromisoformat(day["header"]["date_iso"])
    slug = d.strftime("%d-%m-%Y")
    weekday = day["header"]["weekday"]

    std_name = f"Residents_{slug}.docx"
    veg_name = f"Residents_{slug}_vegan.docx"
    all_name = f"Allergens_Residents_{slug}.docx"
    zip_name = f"{weekday}-{slug}-menus-and-allergens.zip"
    zip_path = os.path.join(out_dir, zip_name)

    # Render to bytes
    std_bytes = render_docxtpl_to_bytes(standard_tpl, std_ctx)
    veg_bytes = render_docxtpl_to_bytes(vegan_tpl, veg_ctx)
    all_doc   = render_allergens_doc(day, std_ctx, veg_ctx, allergens_tpl)
    all_bytes = save_doc_to_bytes(all_doc)

    # Write ZIP only
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(std_name, std_bytes)
        zf.writestr(veg_name, veg_bytes)
        zf.writestr(all_name, all_bytes)

    return zip_path

# ──────────────────────────────────────────────────────────────────────────────
# Template path resolution
# ──────────────────────────────────────────────────────────────────────────────

def resolve_templates(args) -> tuple[str, str, str]:
    """Return (standard_tpl, vegan_tpl, allergens_tpl). Prefer explicit args; else from --templates dir; else ./templates."""
    # explicit wins
    if args.standard_tpl and args.vegan_tpl and args.allergens_tpl:
        return args.standard_tpl, args.vegan_tpl, args.allergens_tpl

    base = args.templates or "templates"
    std = args.standard_tpl or os.path.join(base, "standard.docx")
    veg = args.vegan_tpl or os.path.join(base, "vegan.docx")
    allg= args.allergens_tpl or os.path.join(base, "allergens.docx")
    for p in (std, veg, allg):
        if not os.path.isfile(p):
            raise SystemExit(f"Template not found: {p}")
    return std, veg, allg

# ──────────────────────────────────────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────────────────────────────────────

def main():
    p = argparse.ArgumentParser(description="Generate Henbrook daily menus (ZIP only).")
    p.add_argument("--weekly", required=True, help="Weekly grid DOCX.")
    p.add_argument("--date", help="Generate a single date (YYYY-MM-DD).")
    p.add_argument("--all-days", action="store_true", help="Generate all 7 days.")
    p.add_argument("--out", default="build", help="Output folder for ZIPs.")
    # Templates: either explicit or via a templates folder (default ./templates)
    p.add_argument("--templates", help="Folder with standard.docx, vegan.docx, allergens.docx (default: ./templates)")
    p.add_argument("--standard_tpl", help="Path to Standard template DOCX.")
    p.add_argument("--vegan_tpl", help="Path to Vegan template DOCX.")
    p.add_argument("--allergens_tpl", help="Path to Allergens template DOCX.")
    args = p.parse_args()

    std_tpl, veg_tpl, all_tpl = resolve_templates(args)
    week = parse_week(args.weekly)
    days = week["days"]

    if args.all_days:
        made = [render_day_to_zip(d, std_tpl, veg_tpl, all_tpl, args.out) for d in days]
        print("Generated ZIPs:")
        for z in made: print("-", z)
        return

    if not args.date:
        raise SystemExit("Provide --date YYYY-MM-DD or --all-days.")

    target = next((d for d in days if d["header"]["date_iso"] == args.date), None)
    if not target:
        raise SystemExit(f"Date {args.date} not found in weekly grid.")
    z = render_day_to_zip(target, std_tpl, veg_tpl, all_tpl, args.out)
    print(z)

if __name__ == "__main__":
    main()
